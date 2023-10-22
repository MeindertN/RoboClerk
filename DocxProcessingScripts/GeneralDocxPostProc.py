from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
import argparse
import os
import AIDocxAnnotations as ai
import json


def AddTOCToDocument(document, par, num):
    paragraph = par.clear()
    header = paragraph.insert_paragraph_before("\nTable of Contents\n", "Normal")
    header.runs[0].bold = True
    header.runs[0].underline = True
    header.runs[0].font.size = Pt(12)
    document.paragraphs[num + 1].paragraph_format.page_break_before = True
    header.paragraph_format.page_break_before = True
    header.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    fldChar.set(qn('w:dirty'), 'true')
    fldChar.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p


def DeleteLineFromDocument(par):
    p = par._element
    p.getparent().remove(p)
    par._p = par._element = None


def InsertPageBreak(document, par, num):
    paragraph = par.clear()
    if(len(document.paragraphs) <= num + 1 ):
        document.paragraphs[num].paragraph_format.page_break_before = True
    else:
        document.paragraphs[num + 1].paragraph_format.page_break_before = True

def ProcessParagraphs(document):
    reachedTheEnd = False
    while not reachedTheEnd:
        reachedTheEnd = True
        for num, par in enumerate(document.paragraphs):
            if '~' in par.text and par.text[0] == '~':
                if '~TOC' in par.text:
                    AddTOCToDocument(document, par, num)
                    reachedTheEnd = False
                    break
                elif '~PAGEBREAK' in par.text:
                    InsertPageBreak(document, par, num)
                    reachedTheEnd = False
                    break
        for num, par in enumerate(document.paragraphs):
            if '~' in par.text and par.text[0] == '~':
                if '~REMOVEPARAGRAPH' in par.text:
                    DeleteLineFromDocument(par)
                    reachedTheEnd = False
                    break

# this function scans over all tables in the document and merges any cells that have 
# a single key character in them:
# ^ = merge with cell above 
# < = merge with cell to the left
# > = merge with cell to the right
# note that this only works for rectangular sets of cells! 
def ProcessTables(document):
    for table in document.tables:
        for c,column in enumerate(table.columns):
            for r,cell in enumerate(column.cells):
                if cell.text == '^' and r != 0:
                    cell.text = ''
                    table.cell(r-1,c).merge(table.cell(r,c))
                if cell.text == '<' and c != 0:
                    cell.text = ''
                    table.cell(r,c-1).merge(table.cell(r,c))
                if cell.text == '>' and c < len(table.columns)-1:
                    cell.text = ''
                    table.cell(r,c+1).merge(table.cell(r,c))

def ProcessAIComments(inputFile,document):
    directory, fileWithExtension = os.path.split(inputFile)
    fn, fileExtension = os.path.splitext(fileWithExtension)
    commentFilename = ''
    for filename in os.listdir(directory):
        if filename.endswith("_AIComments.json") and fn.startswith(filename[:-16]):
            commentFilename = filename
    if commentFilename == '':
        return
    commentFilename = os.path.join(directory,commentFilename)
    if os.path.exists(commentFilename):
        with open(commentFilename, 'r') as file:
            comments = [json.loads(line) for line in file]
        ai.ProcessComments(document,comments)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Processes the docx and replaces postprocessing tags (marked with ~ in the"
                                             " document) with elements of the Docx.")

    parser.add_argument('input', type=str, help='the input docx file')
    parser.add_argument('output', type=str, help='the output docx file')
    args = parser.parse_args()

    inputFile = args.input
    outputFile = args.output

    document = Document(inputFile)

    ProcessParagraphs(document)
    ProcessTables(document)
    ProcessAIComments(inputFile,document)
            
    document.save(outputFile)
