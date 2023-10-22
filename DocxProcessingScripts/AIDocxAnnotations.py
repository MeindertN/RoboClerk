import sys
from docx import Document
import argparse
import json

def ProcessComments(doc, comments):
    for comment in comments:
        ScanTables(doc,comment)

def ScanTables(doc, comment):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if ' w:name="comment_'+comment['ID']+'"/>' in paragraph._element.xml:
                        paragraph.add_comment(comment["CommentContent"], author="RoboClerk")
                        return True
    return False # comment not found in tables

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Processes a docx and inserts any AI generated comments into it.")

    parser.add_argument('input', type=str, help='the input docx file')
    parser.add_argument('comments', type=str, help='the file containing the comments')
    parser.add_argument('output', type=str, help='the output docx file')
    args = parser.parse_args()

    inputFile = args.input
    commentFile = args.comments
    outputFile = args.output

    document = Document(inputFile)
    with open(commentFile, 'r') as file:
        comments = [json.loads(line) for line in file]

    ProcessComments(document, comments)
                
    document.save(outputFile)
